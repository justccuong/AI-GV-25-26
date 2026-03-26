from __future__ import annotations

from pathlib import Path
from typing import Iterable

from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, StyleSheet1, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parent
HTML_PATH = ROOT / "Ban_mo_ta_san_pham_EduMind_AI.html"
DOCX_PATH = ROOT / "Ban_mo_ta_san_pham_EduMind_AI.docx"
PDF_PATH = ROOT / "Ban_mo_ta_san_pham_EduMind_AI.pdf"

TIMES_REGULAR = r"C:\Windows\Fonts\times.ttf"
TIMES_BOLD = r"C:\Windows\Fonts\timesbd.ttf"
TIMES_ITALIC = r"C:\Windows\Fonts\timesi.ttf"
TIMES_BOLD_ITALIC = r"C:\Windows\Fonts\timesbi.ttf"


def element_children(node: Tag) -> Iterable[Tag]:
    for child in node.children:
        if isinstance(child, Tag):
            yield child


def text_of(node: Tag) -> str:
    return " ".join(node.get_text(" ", strip=True).split())


def add_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_run_font(run, *, name="Times New Roman", size=13, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_doc_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    style_specs = {
        "CoverEyebrow": {"size": 12, "bold": True},
        "CoverTitle": {"size": 24, "bold": True},
        "CoverSubtitle": {"size": 17, "bold": True},
        "SectionHeading": {"size": 16, "bold": True},
        "SubHeading": {"size": 14, "bold": True},
        "Caption": {"size": 11, "italic": True},
    }

    for name, spec in style_specs.items():
        if name in document.styles:
            style = document.styles[name]
        else:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(spec["size"])
        style.font.bold = spec.get("bold", False)
        style.font.italic = spec.get("italic", False)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(6)


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.0)


def add_paragraph(document: Document, text: str, *, style="Normal", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_cm=1.0) -> None:
    paragraph = document.add_paragraph(style=style)
    paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.first_line_indent = Cm(first_line_cm) if first_line_cm else None
    run = paragraph.add_run(text)
    set_run_font(run, size=13 if style == "Normal" else 12)


def build_docx() -> None:
    soup = BeautifulSoup(HTML_PATH.read_text(encoding="utf-8"), "html.parser")
    document = Document()
    configure_doc_styles(document)
    set_page_layout(document)

    body = soup.body
    assert body is not None

    for node in element_children(body):
        classes = node.get("class", [])

        if node.name == "div" and "cover" in classes:
            for child in element_children(node):
                if child.name == "div" and "eyebrow" in child.get("class", []):
                    paragraph = document.add_paragraph(style="CoverEyebrow")
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(text_of(child))
                    set_run_font(run, size=12, bold=True)
                elif child.name == "h1":
                    paragraph = document.add_paragraph(style="CoverTitle")
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(text_of(child))
                    set_run_font(run, size=24, bold=True)
                elif child.name == "h2":
                    paragraph = document.add_paragraph(style="CoverSubtitle")
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(text_of(child))
                    set_run_font(run, size=17, bold=True)
                elif child.name == "table":
                    rows = child.find_all("tr")
                    table = document.add_table(rows=len(rows), cols=2)
                    table.style = "Table Grid"
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for row_idx, row in enumerate(rows):
                        cells = row.find_all(["td", "th"])
                        for col_idx, cell in enumerate(cells[:2]):
                            target = table.cell(row_idx, col_idx)
                            target.text = text_of(cell)
                            paragraph = target.paragraphs[0]
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            paragraph.paragraph_format.space_after = Pt(0)
                            paragraph.paragraph_format.first_line_indent = None
                            for run in paragraph.runs:
                                set_run_font(run, size=12, bold=(col_idx == 0))
                            if col_idx == 0:
                                add_cell_shading(target, "F2F2F2")
                elif child.name == "p":
                    add_paragraph(document, text_of(child), alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_cm=1.2)
            document.add_page_break()
            continue

        if node.name == "div" and "page-break" in classes:
            continue

        if node.name == "h2":
            paragraph = document.add_paragraph(style="SectionHeading")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = None
            run = paragraph.add_run(text_of(node))
            set_run_font(run, size=16, bold=True)
            continue

        if node.name == "h3":
            paragraph = document.add_paragraph(style="SubHeading")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = None
            run = paragraph.add_run(text_of(node))
            set_run_font(run, size=14, bold=True)
            continue

        if node.name == "p":
            no_indent = "no-indent" in classes
            add_paragraph(
                document,
                text_of(node),
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_cm=0 if no_indent else 1.0,
            )
            continue

        if node.name in {"ol", "ul"}:
            for index, item in enumerate(node.find_all("li", recursive=False), start=1):
                prefix = f"{index}. " if node.name == "ol" else "- "
                paragraph = document.add_paragraph(style="Normal")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.paragraph_format.left_indent = Cm(1.0)
                paragraph.paragraph_format.first_line_indent = Cm(-0.6)
                paragraph.paragraph_format.space_after = Pt(4)
                run = paragraph.add_run(prefix + text_of(item))
                set_run_font(run, size=13)
            continue

        if node.name == "table":
            rows = node.find_all("tr")
            max_cols = max(len(row.find_all(["td", "th"])) for row in rows)
            table = document.add_table(rows=len(rows), cols=max_cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for row_idx, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                for col_idx, cell in enumerate(cells):
                    target = table.cell(row_idx, col_idx)
                    target.text = text_of(cell)
                    paragraph = target.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.first_line_indent = None
                    for run in paragraph.runs:
                        set_run_font(run, size=12, bold=(row_idx == 0))
                    if row_idx == 0:
                        add_cell_shading(target, "F2F2F2")
            continue

        if node.name == "div" and "caption" in classes:
            paragraph = document.add_paragraph(style="Caption")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None
            run = paragraph.add_run(text_of(node))
            set_run_font(run, size=11, italic=True)

    document.save(DOCX_PATH)


def configure_pdf_fonts() -> None:
    pdfmetrics.registerFont(TTFont("TimesNewRoman", TIMES_REGULAR))
    pdfmetrics.registerFont(TTFont("TimesNewRoman-Bold", TIMES_BOLD))
    pdfmetrics.registerFont(TTFont("TimesNewRoman-Italic", TIMES_ITALIC))
    pdfmetrics.registerFont(TTFont("TimesNewRoman-BoldItalic", TIMES_BOLD_ITALIC))


def pdf_styles() -> StyleSheet1:
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="VNNormal",
            parent=styles["Normal"],
            fontName="TimesNewRoman",
            fontSize=13,
            leading=20,
            alignment=TA_JUSTIFY,
            spaceAfter=6,
            firstLineIndent=1.0 * cm,
        )
    )
    styles.add(
        ParagraphStyle(
            name="VNCoverEyebrow",
            parent=styles["Normal"],
            fontName="TimesNewRoman-Bold",
            fontSize=12,
            leading=16,
            alignment=TA_CENTER,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="VNCoverTitle",
            parent=styles["Normal"],
            fontName="TimesNewRoman-Bold",
            fontSize=24,
            leading=30,
            alignment=TA_CENTER,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="VNCoverSubtitle",
            parent=styles["Normal"],
            fontName="TimesNewRoman-Bold",
            fontSize=17,
            leading=24,
            alignment=TA_CENTER,
            spaceAfter=16,
        )
    )
    styles.add(
        ParagraphStyle(
            name="VNSection",
            parent=styles["Normal"],
            fontName="TimesNewRoman-Bold",
            fontSize=16,
            leading=22,
            alignment=TA_LEFT,
            spaceBefore=8,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="VNSubHeading",
            parent=styles["Normal"],
            fontName="TimesNewRoman-Bold",
            fontSize=14,
            leading=19,
            alignment=TA_LEFT,
            spaceBefore=6,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="VNCaption",
            parent=styles["Normal"],
            fontName="TimesNewRoman-Italic",
            fontSize=11,
            leading=14,
            alignment=TA_CENTER,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="VNList",
            parent=styles["VNNormal"],
            leftIndent=1.0 * cm,
            firstLineIndent=-0.6 * cm,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="VNNoIndent",
            parent=styles["VNNormal"],
            firstLineIndent=0,
        )
    )
    return styles


def build_pdf() -> None:
    configure_pdf_fonts()
    styles = pdf_styles()
    soup = BeautifulSoup(HTML_PATH.read_text(encoding="utf-8"), "html.parser")
    body = soup.body
    assert body is not None

    story = []

    def add_table(node: Tag, first_col_gray=False) -> None:
        rows = []
        for row in node.find_all("tr", recursive=False):
            rows.append([Paragraph(text_of(cell), styles["VNNoIndent"]) for cell in row.find_all(["td", "th"], recursive=False)])
        table = Table(rows, repeatRows=1)
        table_style = TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.8, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F2F2")),
                ("FONTNAME", (0, 0), (-1, 0), "TimesNewRoman-Bold"),
                ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
        if first_col_gray:
            table_style.add("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F2F2F2"))
            table_style.add("FONTNAME", (0, 0), (0, -1), "TimesNewRoman-Bold")
        table.setStyle(table_style)
        story.append(table)
        story.append(Spacer(1, 0.2 * cm))

    for node in element_children(body):
        classes = node.get("class", [])

        if node.name == "div" and "cover" in classes:
            for child in element_children(node):
                if child.name == "div" and "eyebrow" in child.get("class", []):
                    story.append(Spacer(1, 3.2 * cm))
                    story.append(Paragraph(text_of(child), styles["VNCoverEyebrow"]))
                elif child.name == "h1":
                    story.append(Paragraph(text_of(child), styles["VNCoverTitle"]))
                elif child.name == "h2":
                    story.append(Paragraph(text_of(child), styles["VNCoverSubtitle"]))
                elif child.name == "table":
                    add_table(child, first_col_gray=True)
                elif child.name == "p":
                    story.append(Paragraph(text_of(child), styles["VNNormal"]))
            story.append(PageBreak())
            continue

        if node.name == "div" and "page-break" in classes:
            continue

        if node.name == "h2":
            story.append(Paragraph(text_of(node), styles["VNSection"]))
            continue

        if node.name == "h3":
            story.append(Paragraph(text_of(node), styles["VNSubHeading"]))
            continue

        if node.name == "p":
            style_name = "VNNoIndent" if "no-indent" in classes else "VNNormal"
            story.append(Paragraph(text_of(node), styles[style_name]))
            continue

        if node.name in {"ol", "ul"}:
            for index, item in enumerate(node.find_all("li", recursive=False), start=1):
                prefix = f"{index}. " if node.name == "ol" else "- "
                story.append(Paragraph(prefix + text_of(item), styles["VNList"]))
            continue

        if node.name == "table":
            add_table(node)
            continue

        if node.name == "div" and "caption" in classes:
            story.append(Paragraph(text_of(node), styles["VNCaption"]))

    document = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=2.6 * cm,
        rightMargin=2.0 * cm,
        topMargin=2.2 * cm,
        bottomMargin=2.2 * cm,
    )
    document.build(story)


def main() -> None:
    build_docx()
    build_pdf()
    print(f"DOCX={DOCX_PATH.name}")
    print(f"PDF={PDF_PATH.name}")


if __name__ == "__main__":
    main()
